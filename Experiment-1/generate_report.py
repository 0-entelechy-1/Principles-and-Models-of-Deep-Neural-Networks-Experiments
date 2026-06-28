import os
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'outputs')
CODE_FILE = os.path.join(BASE_DIR, 'experiment1.py')
OUTPUT_PATH = os.path.join(BASE_DIR, 'docs', '实验总结一.docx')
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)


def set_font(run, name='宋体', size=Pt(12), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def para(doc, text, bold=False, indent=False, font='宋体', size=Pt(12)):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, font, size, bold)
    return p


def img(doc, path, caption, width=Cm(15)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    set_font(r, '宋体', Pt(10), bold=True)


def code_block(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 一、实验目的与要求 =====
    heading(doc, '一、实验目的与要求', 1)
    para(doc,
         '利用PyTorch深度学习框架，学习与验证神经网络实现回归任务的基本运算原理，'
         '理解多层感知机在曲线拟合的回归任务建模、训练模型、评估模型等方法；'
         '理解梯度下降及常见的改进方法、反向传播算法原理、梯度消失与梯度爆炸的机理'
         '以及使用相应方法缓解这些问题的策略。', indent=True)

    # ===== 二、实验原理 =====
    heading(doc, '二、实验原理', 1)

    heading(doc, '2.1 多层感知机（MLP）结构与运算原理', 2)
    para(doc,
         '多层感知机（MLP）由输入层、若干隐藏层和输出层组成。每一层包含一组神经元，'
         '相邻层之间通过权重矩阵全连接。对于第l层，其前向传播公式为：', indent=True)
    para(doc, 'z^(l) = W^(l) · a^(l-1) + b^(l)', font='Courier New', size=Pt(11))
    para(doc, 'a^(l) = σ(z^(l))', font='Courier New', size=Pt(11))
    para(doc,
         '其中W^(l)为权重矩阵，b^(l)为偏置向量，σ为激活函数。本实验使用了ReLU、'
         'Sigmoid和Tanh三种激活函数。ReLU(x)=max(0,x)，计算高效且不易梯度消失；'
         'Sigmoid(x)=1/(1+e^(-x))，输出范围(0,1)，但容易导致梯度消失。', indent=True)

    heading(doc, '2.2 梯度下降及改进方法', 2)
    para(doc,
         '标准梯度下降（SGD）：θ = θ - η∇L(θ)，每次沿负梯度方向更新参数，'
         'η为学习率，控制更新步长。', indent=True)
    para(doc,
         '动量法（Momentum）：v_t = γv_{t-1} + η∇L(θ)，θ = θ - v_t，'
         '利用历史梯度信息加速收敛，减少振荡。γ为动量系数，通常取0.9。', indent=True)
    para(doc,
         'Adam优化器：结合了动量法和RMSProp的优点，维护一阶矩估计和二阶矩估计，'
         '自适应调整每个参数的学习率，在大多数场景下收敛速度快且稳定。', indent=True)

    heading(doc, '2.3 反向传播算法', 2)
    para(doc,
         '反向传播算法通过链式法则计算损失函数对每个参数的梯度。从输出层开始，'
         '逐层向前计算误差项：δ^(l) = (W^(l+1))^T · δ^(l+1) ⊙ σ\'(z^(l))，'
         '然后计算参数梯度 ∂L/∂W^(l) = δ^(l)·(a^(l-1))^T。'
         '对于MSE损失 L = (1/N)Σ(y_i - ŷ_i)^2，输出层误差 δ^(L) = -(2/N)(y - ŷ)。', indent=True)

    heading(doc, '2.4 梯度消失与梯度爆炸', 2)
    para(doc,
         '梯度消失：当使用Sigmoid等饱和激活函数时，其导数最大仅为0.25，'
         '反向传播中梯度逐层相乘，靠近输入层的梯度会指数级衰减，导致浅层参数几乎不更新。', indent=True)
    para(doc,
         '梯度爆炸：当学习率过大或权重初始化不当时，梯度可能指数级增长，'
         '导致参数更新过大，损失变为NaN。', indent=True)
    para(doc,
         '缓解方法：使用ReLU激活函数（导数为0或1）、合适的权重初始化（如Xavier/He初始化）、'
         '梯度裁剪、BatchNorm等。', indent=True)

    # ===== 三、实验环境与数据集 =====
    heading(doc, '三、实验环境与数据集', 1)

    heading(doc, '3.1 实验环境', 2)
    para(doc, '编程语言：Python 3.9', indent=True)
    para(doc, '深度学习框架：PyTorch 2.7.0', indent=True)
    para(doc, '运行设备：CPU', indent=True)
    para(doc, '依赖库：numpy, matplotlib', indent=True)

    heading(doc, '3.2 数据集构建', 2)
    para(doc,
         '目标函数：y = sin(x) + 0.5·cos(2x)，x ∈ [-π, π]', indent=True)
    para(doc,
         '添加高斯噪声（标准差0.1），模拟真实数据中的测量误差，生成300个样本。', indent=True)
    para(doc,
         '数据集划分：训练集70%（210个）、验证集15%（45个）、测试集15%（45个）。'
         '验证集用于训练过程中的模型选择，测试集用于最终性能评估。', indent=True)

    img(doc, os.path.join(IMG_DIR, 'dataset.png'), '图1 数据集分布')

    # ===== 四、模型构建 =====
    heading(doc, '四、模型构建', 1)

    heading(doc, '4.1 MLP模型结构', 2)
    para(doc,
         '实现了通用的MLP类，支持自定义隐藏层大小和激活函数。实验设计了4种网络结构：', indent=True)
    para(doc, '浅层网络(1层64)：1个隐藏层，64个神经元，ReLU激活，参数量193', indent=True)
    para(doc, '中层网络(2层64)：2个隐藏层，各64个神经元，ReLU激活，参数量4353', indent=True)
    para(doc, '深层网络(4层64)：4个隐藏层，各64个神经元，ReLU激活，参数量12673', indent=True)
    para(doc, '深层Sigmoid(4层64)：4个隐藏层，各64个神经元，Sigmoid激活，参数量12673', indent=True)

    heading(doc, '4.2 训练配置', 2)
    para(doc, '损失函数：MSE（均方误差）', indent=True)
    para(doc, '批大小：32', indent=True)
    para(doc, '训练轮数：500（对比实验3为300）', indent=True)
    para(doc, '权重初始化：PyTorch默认（Kaiming均匀初始化）', indent=True)

    # ===== 五、实验结果与分析 =====
    heading(doc, '五、实验结果与分析', 1)

    heading(doc, '5.1 对比实验1：不同优化器对比', 2)
    para(doc,
         '使用中层网络(2层64, ReLU)，学习率0.01，训练500轮，'
         '对比SGD、SGD+Momentum、Adam三种优化器的训练效果。', indent=True)

    heading(doc, '实验结果：', 3)
    para(doc, 'SGD：收敛较慢，最终验证MSE = 0.013965', indent=True)
    para(doc, 'SGD+Momentum：收敛更快更稳，最终验证MSE = 0.012091', indent=True)
    para(doc, 'Adam：初期收敛最快，最终验证MSE = 0.014744', indent=True)

    img(doc, os.path.join(IMG_DIR, 'optimizer_comparison.png'), '图2 不同优化器的训练/验证损失对比')
    img(doc, os.path.join(IMG_DIR, 'optimizer_fitting.png'), '图3 不同优化器的曲线拟合结果')

    heading(doc, '分析：', 3)
    para(doc,
         'SGD+Momentum通过引入动量项（γ=0.9），在标准SGD基础上加速了收敛，'
         '同时保持了较好的稳定性，取得了最优的最终性能。'
         'Adam结合了自适应学习率和动量，初期收敛最快，但由于学习率自适应调整，'
         '后期存在一定振荡，最终性能略逊于SGD+Momentum。'
         '这说明对于本实验规模的简单回归任务，SGD+Momentum是更优的选择。', indent=True)

    heading(doc, '5.2 对比实验2：不同网络结构与激活函数对比', 2)
    para(doc, '使用Adam优化器，学习率0.01，训练500轮。', indent=True)

    heading(doc, '实验结果：', 3)
    para(doc, '浅层网络(1层64, ReLU)：测试MSE = 0.009497', indent=True)
    para(doc, '中层网络(2层64, ReLU)：测试MSE = 0.008861（最佳）', indent=True)
    para(doc, '深层网络(4层64, ReLU)：测试MSE = 0.009183', indent=True)
    para(doc, '深层Sigmoid(4层64, Sigmoid)：测试MSE = 0.010886', indent=True)

    img(doc, os.path.join(IMG_DIR, 'structure_comparison.png'), '图4 不同网络结构的训练/验证损失对比')
    img(doc, os.path.join(IMG_DIR, 'structure_fitting.png'), '图5 不同网络结构的曲线拟合结果')

    heading(doc, '分析：', 3)
    para(doc,
         '中层网络(2层64)取得了最佳测试性能（MSE=0.008861）。浅层网络（1层64）'
         '表达能力相对不足；深层网络（4层64）虽然参数更多，但容易过拟合且训练不稳定，'
         '验证损失波动较大。深层Sigmoid网络由于梯度消失问题，性能最差（MSE=0.010886），'
         '验证了Sigmoid在深层网络中的局限性。', indent=True)
    para(doc,
         '这一结果说明网络深度并非越深越好，需要根据任务复杂度选择合适的网络结构。'
         '对于本实验中sin(x)+0.5cos(2x)这样的简单函数拟合，2层隐藏层已经足够。', indent=True)

    heading(doc, '5.3 对比实验3：不同学习率对比', 2)
    para(doc,
         '使用中层网络(2层64, ReLU)，SGD优化器，训练300轮，'
         '对比学习率0.001、0.01、0.1、1.0。', indent=True)

    heading(doc, '实验结果：', 3)
    para(doc, 'lr=0.001：收敛极慢，最终验证MSE = 0.206142', indent=True)
    para(doc, 'lr=0.01：收敛适中，最终验证MSE = 0.014782', indent=True)
    para(doc, 'lr=0.1：收敛较快，最终验证MSE = 0.022413', indent=True)
    para(doc, 'lr=1.0：梯度爆炸，损失变为NaN', indent=True)

    img(doc, os.path.join(IMG_DIR, 'lr_comparison.png'), '图6 不同学习率的训练/验证损失对比')

    heading(doc, '分析：', 3)
    para(doc,
         '学习率是最关键的超参数之一。lr=0.001时收敛极慢，300轮后仍未充分学习；'
         'lr=0.01在SGD下取得了较好的平衡，收敛稳定且最终性能较好；'
         'lr=0.1虽然初期收敛更快，但由于步长偏大导致后期振荡；'
         'lr=1.0时SGD出现NaN损失，这是典型的梯度爆炸现象。'
         '实际应用中常使用学习率调度策略（如余弦退火、阶梯衰减）来兼顾快速收敛和稳定训练。', indent=True)

    # ===== 六、梯度消失与梯度爆炸分析 =====
    heading(doc, '六、梯度消失与梯度爆炸分析', 1)

    heading(doc, '6.1 各模型梯度范数变化', 2)
    img(doc, os.path.join(IMG_DIR, 'gradient_analysis.png'), '图7 各模型各层梯度范数变化')
    para(doc,
         '从梯度范数图可以看出，使用ReLU激活函数的网络各层梯度范数较为均匀，'
         '没有明显的梯度消失或爆炸现象。这是因为ReLU的导数在正区间为1，'
         '梯度在反向传播中不会大幅衰减。', indent=True)
    para(doc,
         '而使用Sigmoid激活函数的深层网络，靠近输入层的梯度范数明显小于输出层，'
         '呈现典型的梯度消失趋势。Sigmoid函数的导数最大值仅为0.25，'
         '多个小于1的导数值连乘导致梯度指数级衰减。', indent=True)

    heading(doc, '6.2 深层Sigmoid网络梯度消失详细分析', 2)
    para(doc,
         '对6层Sigmoid网络（6个隐藏层各64个神经元，参数量20993）进行200轮训练，'
         '追踪各层权重梯度范数。', indent=True)
    img(doc, os.path.join(IMG_DIR, 'gradient_vanishing.png'), '图8 6层Sigmoid网络梯度消失现象')
    para(doc,
         '6层Sigmoid网络中，第1层（最靠近输入层）的梯度范数远小于第6层（最靠近输出层），'
         '差距可达数个数量级。这是典型的梯度消失现象。'
         'Sigmoid函数的导数σ\'(x)=σ(x)(1-σ(x))，最大值仅为0.25，'
         '在反向传播中多个小于1的值相乘，导致浅层梯度指数级衰减，'
         '浅层参数几乎得不到有效更新。', indent=True)

    heading(doc, '6.3 梯度爆炸现象', 2)
    para(doc,
         '在对比实验3中，学习率lr=1.0时SGD出现NaN损失，这是典型的梯度爆炸现象。'
         '过大的学习率使得参数更新步长过大，损失函数值剧烈振荡直至数值溢出。'
         '可以通过降低学习率、使用梯度裁剪（torch.nn.utils.clip_grad_norm_）'
         '或采用学习率预热策略来缓解。', indent=True)

    # ===== 七、最佳模型评估 =====
    heading(doc, '七、最佳模型评估', 1)
    para(doc,
         '综合所有对比实验，最佳模型为中层网络(2层64, ReLU, Adam)，'
         '测试MSE = 0.008861。该模型在表达能力和训练稳定性之间取得了良好平衡。', indent=True)
    img(doc, os.path.join(IMG_DIR, 'best_model_fitting.png'), '图9 最佳模型曲线拟合结果')
    para(doc,
         '从拟合结果图可以看出，最佳模型能够很好地捕捉目标函数sin(x)+0.5cos(2x)的整体趋势，'
         '拟合曲线与真实函数曲线高度重合，残差主要集中在噪声较大的区域。', indent=True)

    # ===== 八、总结与心得体会 =====
    heading(doc, '八、总结与心得体会', 1)
    para(doc,
         '通过本次实验，我深入理解了多层感知机在曲线拟合回归任务中的应用，主要收获如下：', indent=True)
    para(doc,
         '1. 网络结构选择：并非越深越好。对于本实验中的简单曲线拟合任务，'
         '2层隐藏层已经足够表达目标函数，更深的网络反而可能带来过拟合和训练不稳定问题。'
         '实际项目中应根据任务复杂度选择适当的网络深度。', indent=True)
    para(doc,
         '2. 优化器选择：Adam优化器在大多数情况下收敛最快，但SGD+Momentum在最终性能上可能更优。'
         '选择优化器需要根据具体任务进行实验验证，不能一概而论。', indent=True)
    para(doc,
         '3. 学习率的重要性：学习率是最关键的超参数之一。过小导致收敛慢，过大导致梯度爆炸。'
         '实际应用中常使用学习率衰减策略来兼顾初期快速收敛和后期稳定优化。', indent=True)
    para(doc,
         '4. 梯度消失与爆炸：Sigmoid激活函数在深层网络中容易导致梯度消失，'
         'ReLU是更好的选择。对于梯度爆炸，可以通过梯度裁剪和适当的学习率来缓解。'
         '通过实验直观地观察到了这些现象，加深了对理论知识的理解。', indent=True)
    para(doc,
         '5. 反向传播的理解：通过追踪各层梯度范数，直观地看到了反向传播中梯度的流动情况，'
         '加深了对链式法则和梯度传播机制的理解。这对后续学习更复杂的深度网络（如CNN、RNN）'
         '打下了坚实的理论基础。', indent=True)

    # ===== 附录 =====
    heading(doc, '附录：实验代码', 1)
    with open(CODE_FILE, 'r', encoding='utf-8') as f:
        code_text = f.read()
    code_block(doc, code_text)

    doc.save(OUTPUT_PATH)
    print(f'报告已生成: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
