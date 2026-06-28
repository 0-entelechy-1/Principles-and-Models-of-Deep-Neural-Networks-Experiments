import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'outputs')
CODE_FILE = os.path.join(BASE_DIR, 'experiment2.py')
OUTPUT_PATH = os.path.join(BASE_DIR, 'docs', '实验总结二.docx')
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

LOG_FILE = os.path.join(IMG_DIR, 'experiment_log.txt')
SUPPLEMENT_LOG_FILE = os.path.join(IMG_DIR, 'experiment2_supplement_log.txt')


def set_font(run, name='宋体', size=Pt(12), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def para(doc, text, bold=False, indent=False, font='宋体', size=Pt(12)):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, font, size, bold)


def img(doc, path, caption, width=Cm(15)):
    if not os.path.exists(path):
        para(doc, f'[图片缺失: {caption}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    set_font(r, '宋体', Pt(10), bold=True)


def code_block(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')


def parse_log_file(path, results):
    if not os.path.exists(path):
        return
    import re
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    for m in re.finditer(r'--- 训练 (\S+.*?) ---\n(.*?)(?=---|\Z)', content, re.DOTALL):
        name = m.group(1)
        block = m.group(2)
        acc_m = re.search(r'最终测试准确率: ([\d.]+)', block)
        time_m = re.search(r'耗时: ([\d.]+)s', block)
        results[name] = {
            'acc': float(acc_m.group(1)) if acc_m else 0,
            'time': float(time_m.group(1)) if time_m else 0,
        }


def parse_log():
    results = {}
    parse_log_file(LOG_FILE, results)
    parse_log_file(SUPPLEMENT_LOG_FILE, results)
    return results


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    log_data = parse_log()

    heading(doc, '一、实验目的与要求', 1)
    para(doc,
         '利用PyTorch深度学习框架以及相应的深度神经网络，学习与验证神经网络实现分类任务的基本原理，'
         '掌握构建LeNet、VGGNet、ResNet等经典卷积神经网络模型的基本方法，'
         '理解注意力机制、残差机制及各类卷积的作用，从而实现能够进行复杂分类任务的设计目标。', indent=True)

    heading(doc, '二、实验原理', 1)

    heading(doc, '2.1 卷积神经网络（CNN）基本构成', 2)
    para(doc,
         '卷积神经网络是深度学习中处理图像数据的核心架构，其基本构成要素包括：', indent=True)
    para(doc,
         '(1) 卷积层（Convolutional Layer）：通过可学习的卷积核提取局部特征，'
         '具有参数共享和局部连接的特点，大幅减少参数量；', indent=True)
    para(doc,
         '(2) 池化层（Pooling Layer）：对特征图进行下采样，减小空间尺寸，增大感受野，'
         '常用的有最大池化和平均池化；', indent=True)
    para(doc,
         '(3) 批归一化（Batch Normalization）：对每个小批量的 activations 进行标准化，'
         '加速训练并提高模型稳定性；', indent=True)
    para(doc,
         '(4) 激活函数（Activation Function）：引入非线性变换，ReLU是最常用的选择；', indent=True)
    para(doc,
         '(5) 全连接层（Fully Connected Layer）：在分类任务中通常位于网络末端，'
         '将特征映射到类别空间。', indent=True)

    heading(doc, '2.2 残差结构与ResNet', 2)
    para(doc,
         'ResNet（Residual Network）由He等（2015）提出，通过引入残差连接（Skip Connection）'
         '解决了深层网络的退化问题。残差块的输出为：y = F(x) + x，'
         '其中F(x)是两层卷积的输出。这使得网络学习的是残差映射F(x)而非完整映射，'
         '梯度可以通过跳跃连接直接传播到浅层，有效缓解了梯度消失问题。'
         'ResNet-18包含4个残差阶段，通道数依次为64、128、256、512。', indent=True)

    heading(doc, '2.3 注意力机制', 2)
    para(doc,
         '注意力机制通过自适应地调整特征权重，增强重要特征、抑制无关特征。'
         '本实验验证了以下注意力机制：', indent=True)
    para(doc,
         '(1) SE（Squeeze-and-Excitation）模块：通道域注意力，先通过全局平均池化压缩空间信息（Squeeze），'
         '再通过两层全连接网络学习通道间关系（Excitation），最后通过Sigmoid生成通道权重。'
         'SE模块仅增加极少的参数（约2/r倍通道数），却能显著提升特征表示能力；', indent=True)
    para(doc,
         '(2) CBAM（Convolutional Block Attention Module）：混合域注意力，'
         '结合通道注意力和空间注意力，先对通道维度加权，再对空间位置加权；', indent=True)
    para(doc,
         '(3) 自注意力机制（Self-Attention）：通过Query-Key-Value计算特征内部的全局依赖关系，'
         '在Vision Transformer中广泛应用。', indent=True)

    heading(doc, '2.4 各类卷积', 2)
    para(doc,
         '(1) 标准卷积：C_in×C_out×K×K个参数，计算量为H×W×C_in×C_out×K×K；', indent=True)
    para(doc,
         '(2) 分组卷积（Grouped Convolution）：将输入通道分成G组，每组独立卷积，参数量减少为1/G；', indent=True)
    para(doc,
         '(3) 逐点卷积（Pointwise Convolution）：1×1卷积，用于通道间特征融合和通道数变换；', indent=True)
    para(doc,
         '(4) 深度可分离卷积（Depthwise Separable Convolution）：先对每个通道独立做深度卷积（K×K, groups=C_in），'
         '再做1×1逐点卷积。参数量从C_in×C_out×K×K降为C_in×K×K+C_in×C_out，约减少8-9倍；', indent=True)
    para(doc,
         '(5) 空洞卷积（Dilated Convolution）：在卷积核元素间插入空洞（膨胀率d），'
         '感受野从K扩大到K+(K-1)(d-1)，不增加参数量；', indent=True)
    para(doc,
         '(6) 转置卷积（Transposed Convolution）：用于上采样，通过反向的卷积操作增大特征图尺寸。', indent=True)

    heading(doc, '三、实验环境与数据集', 1)

    heading(doc, '3.1 实验环境', 2)
    para(doc, '编程语言：Python 3.9', indent=True)
    para(doc, '深度学习框架：PyTorch 2.7.0, torchvision 0.22.0', indent=True)
    para(doc, '运行设备：CPU', indent=True)
    para(doc, '依赖库：numpy, matplotlib, Pillow', indent=True)

    heading(doc, '3.2 CIFAR-10数据集', 2)
    para(doc,
         'CIFAR-10数据集包含60000张32×32彩色图像，分为10个类别（飞机、汽车、鸟、猫、鹿、'
         '狗、青蛙、马、船、卡车），其中训练集50000张，测试集10000张。'
         '训练时使用了数据增强（随机水平翻转和随机裁剪），测试时仅做标准化。', indent=True)
    img(doc, os.path.join(IMG_DIR, 'dataset_samples.png'), '图1 CIFAR-10数据集样本')

    heading(doc, '四、模型构建', 1)

    heading(doc, '4.1 MLP基线模型', 2)
    para(doc,
         '将图像展平为一维向量（3×32×32=3072），通过3层全连接网络（3072→512→256→10），'
         '使用ReLU激活和Dropout正则化。该模型作为基线，用于对比CNN的空间特征提取能力。', indent=True)

    heading(doc, '4.2 VGG-style模型', 2)
    para(doc,
         '自定义的VGG风格网络，包含3个卷积块（每块2层卷积+BN+ReLU+MaxPool），'
         '通道数64→128→256，分类头为4096→512→10。', indent=True)

    heading(doc, '4.3 ResNet-18', 2)
    para(doc,
         '实现了标准ResNet-18结构，包含4个残差阶段，使用BasicBlock（两层3×3卷积+残差连接），'
         '最终通过全局平均池化和全连接层输出分类结果。', indent=True)

    heading(doc, '4.4 注意力机制模型', 2)
    para(doc,
         '在相同的基础CNN架构上，分别实现无注意力模块和带SE模块的两个版本，'
         '用于验证SE注意力机制对分类性能的影响。', indent=True)

    heading(doc, '4.5 卷积类型模型', 2)
    para(doc,
         '在相同的基础架构上，分别使用标准卷积和深度可分离卷积，'
         '对比参数量和分类性能的差异。', indent=True)

    heading(doc, '4.6 MobileNetV2 风格轻量模型', 2)
    para(doc,
         'MobileNetV2 是面向移动端设计的轻量化网络，核心为倒置残差块（Inverted Residual）：'
         '先通过 1×1 卷积升维，再做 3×3 深度可分离卷积，最后通过 1×1 卷积降维，'
         '并在输入输出通道相同且步长为 1 时引入残差连接。'
         '本实验在其基础上分别构建无 SE 注意力与带 SE 注意力的版本，以验证注意力机制在轻量模型中的作用。', indent=True)

    heading(doc, '4.7 ShuffleNetV2 风格轻量模型', 2)
    para(doc,
         'ShuffleNetV2 是另一种高效的轻量化网络，通过通道拆分（Channel Split）与通道洗牌（Channel Shuffle）'
         '促进不同分组间的信息交流，同时大量使用 1×1 卷积、3×3 分组卷积和深度可分离卷积降低计算量。'
         '本实验分别构建关闭与开启 Channel Shuffle 的两个版本，以验证通道洗牌对分组卷积特征融合的影响。', indent=True)

    heading(doc, '五、实验结果与分析', 1)

    heading(doc, '5.1 主对比实验：MLP vs VGG-style vs ResNet-18', 2)
    img(doc, os.path.join(IMG_DIR, 'main_comparison.png'), '图2 MLP vs VGG-style vs ResNet-18 训练曲线对比')
    img(doc, os.path.join(IMG_DIR, 'main_bar_comparison.png'), '图3 最终准确率与参数量对比')

    if log_data:
        for name, data in log_data.items():
            if name in ['MLP', 'VGG-style', 'ResNet-18']:
                para(doc, f'{name}：最终测试准确率 = {data["acc"]:.4f}，训练耗时 = {data["time"]:.1f}s', indent=True)

    para(doc,
         'ResNet-18凭借残差连接在深层网络中保持了良好的梯度传播，取得了最高的测试准确率。'
         'VGG-style模型虽然参数量较大，但缺乏残差连接，在20个epoch内收敛不如ResNet充分。'
         'MLP基线模型忽略了图像的空间结构信息，将像素展平后处理，性能显著低于两种CNN模型，'
         '验证了卷积操作在图像任务中的关键作用。', indent=True)

    heading(doc, '5.2 预测结果可视化', 2)
    img(doc, os.path.join(IMG_DIR, 'prediction_samples.png'), '图4 各模型预测结果示例')
    para(doc,
         '从预测结果示例可以看出，CNN模型（尤其是ResNet-18）在大多数类别上都能做出正确预测，'
         '而MLP模型由于缺乏空间特征提取能力，错误率明显更高。'
         'CNN容易识别的类别包括形状特征明显的汽车、卡车等，'
         '较难识别的类别包括外观相似的猫和狗、鸟等。', indent=True)

    heading(doc, '5.3 注意力机制验证', 2)
    img(doc, os.path.join(IMG_DIR, 'attention_comparison.png'), '图5 SE注意力机制对比')
    para(doc,
         '在相同架构上添加SE模块后，模型能够自适应地增强重要通道的特征响应，'
         '抑制无关通道的噪声。SE模块仅增加极少的额外参数（约每个模块2×C/r个），'
         '但通过通道间的信息交互提升了特征表示的判别力。'
         '实验结果表明，带SE模块的模型在收敛速度和最终准确率上均优于无注意力的基线。', indent=True)

    heading(doc, '5.4 卷积类型验证', 2)
    img(doc, os.path.join(IMG_DIR, 'conv_type_comparison.png'), '图6 标准卷积 vs 深度可分离卷积对比')
    para(doc,
         '深度可分离卷积将标准卷积分解为深度卷积和逐点卷积两步，'
         '参数量大幅减少（约为标准卷积的1/8至1/9）。实验结果显示，'
         '深度可分离卷积在参数量显著降低的情况下，仍能保持接近标准卷积的分类准确率，'
         '验证了其在轻量化模型设计中的价值。MobileNet系列正是基于深度可分离卷积构建的。', indent=True)

    heading(doc, '5.5 MobileNetV2 注意力机制验证', 2)
    img(doc, os.path.join(IMG_DIR, 'mobilenet_attention_comparison.png'),
        '图7 MobileNetV2 有无 SE 注意力对比')
    img(doc, os.path.join(IMG_DIR, 'mobile_shuffle_bar_comparison.png'),
        '图8 轻量模型消融实验最终准确率')

    if log_data:
        for name, data in log_data.items():
            if name.startswith('MobileNetV2'):
                para(doc, f'{name}：最终测试准确率 = {data["acc"]:.4f}，训练耗时 = {data["time"]:.1f}s', indent=True)

    para(doc,
         '在 MobileNetV2 风格的轻量化网络上引入 SE 注意力模块，模型能够以极小的参数量代价'
         '（仅增加通道注意力全连接层）自适应地增强重要通道响应。实验结果表明，'
         '带 SE 模块的 MobileNetV2 收敛更快、最终准确率更高，'
         '验证了注意力机制在轻量化模型中同样能显著提升特征表达能力。', indent=True)

    heading(doc, '5.6 ShuffleNetV2 特殊卷积验证', 2)
    img(doc, os.path.join(IMG_DIR, 'shufflenet_shuffle_comparison.png'),
        '图9 ShuffleNetV2 有无 Channel Shuffle 对比')

    if log_data:
        for name, data in log_data.items():
            if name.startswith('ShuffleNetV2'):
                para(doc, f'{name}：最终测试准确率 = {data["acc"]:.4f}，训练耗时 = {data["time"]:.1f}s', indent=True)

    para(doc,
         'ShuffleNetV2 通过 Channel Shuffle 将分组卷积产生的特征在通道维度重新排列，'
         '打破不同分组之间的信息隔离，增强特征融合能力。'
         '对比实验显示，开启 Channel Shuffle 的模型准确率明显高于关闭 shuffle 的版本，'
         '说明通道洗牌是 ShuffleNetV2 提升分组卷积表达力的关键设计。', indent=True)

    heading(doc, '5.7 特征图可视化', 2)
    img(doc, os.path.join(IMG_DIR, 'feature_maps.png'), '图10 VGG-style特征图可视化')
    para(doc,
         '通过可视化VGG-style模型各层的特征图，可以观察到：'
         '浅层特征图主要捕捉边缘、纹理等低级特征；'
         '深层特征图则呈现出更加抽象的语义信息，响应特定的目标部件或整体结构。'
         '这种从低级到高级的特征层次是CNN能够进行复杂图像分类的基础。', indent=True)

    heading(doc, '5.6 MLP与CNN对比分析', 2)
    para(doc,
         'MLP与CNN在图像分类任务上的核心差异：', indent=True)
    para(doc,
         '(1) 空间信息：MLP将图像展平，丢失了像素间的空间关系；CNN通过卷积核保持并利用空间局部性；', indent=True)
    para(doc,
         '(2) 参数效率：MLP的全连接层参数量巨大（3072×512=1,572,864）；CNN通过参数共享大幅减少参数；', indent=True)
    para(doc,
         '(3) 平移不变性：CNN的卷积操作天然具有平移不变性，目标在图像中的位置变化不影响特征提取；MLP不具备此特性；', indent=True)
    para(doc,
         '(4) 泛化能力：CNN在图像任务上的泛化能力远强于MLP，这是实验准确率差距的根本原因。', indent=True)

    heading(doc, '六、总结与心得体会', 1)
    para(doc,
         '通过本次图像分类实验，我深入理解了卷积神经网络的设计原理和关键组件，主要收获如下：', indent=True)
    para(doc,
         '1. CNN vs MLP：卷积操作是处理图像数据的关键，CNN通过局部连接和参数共享高效提取空间特征，'
         '性能远超将图像展平处理的MLP基线。', indent=True)
    para(doc,
         '2. 残差连接的价值：ResNet通过跳跃连接解决了深层网络的退化问题，'
         '使得训练更深的网络成为可能。残差学习的思想对后续许多网络架构产生了深远影响。', indent=True)
    para(doc,
         '3. 注意力机制的作用：SE模块以极小的参数代价提升了特征表示能力，'
         '体现了"关注重要信息"这一核心思想。现代网络（如EfficientNet）普遍集成了注意力机制。', indent=True)
    para(doc,
         '4. 轻量化设计：深度可分离卷积通过分解标准卷积大幅减少计算量和参数量，'
         '是MobileNet等轻量化模型的基础，对移动端部署至关重要。', indent=True)
    para(doc,
         '5. MobileNetV2与ShuffleNetV2：MobileNetV2利用倒置残差与深度可分离卷积实现高效特征提取，'
         'ShuffleNetV2则通过Channel Shuffle增强分组卷积的信息流动；'
         '在二者上分别验证注意力机制与特殊卷积的作用，结果与理论预期一致。', indent=True)
    para(doc,
         '6. 数据增强的重要性：随机翻转和裁剪等数据增强手段有效增加了训练数据的多样性，'
         '减轻了过拟合问题，是提升模型泛化能力的实用技巧。', indent=True)

    heading(doc, '附录：实验代码', 1)
    with open(CODE_FILE, 'r', encoding='utf-8') as f:
        code_block(doc, f.read())

    doc.save(OUTPUT_PATH)
    print(f'报告已生成: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
