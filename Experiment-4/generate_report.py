import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, 'outputs')
CODE_FILE = os.path.join(BASE_DIR, 'experiment4.py')
OUTPUT_PATH = os.path.join(BASE_DIR, 'docs', '实验总结四.docx')
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)


def set_font(run, name='宋体', size=Pt(12), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def para(doc, text, bold=False, indent=False, font='宋体', size=Pt(12)):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, font, size, bold)


def img(doc, path, caption, width=Cm(15)):
    if not os.path.exists(path):
        para(doc, f'[图片缺失: {caption}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    set_font(r, '宋体', Pt(10), bold=True)


def code_block(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 一、实验目的与要求 =====
    heading(doc, '一、实验目的与要求', 1)
    para(doc,
         '利用PyTorch深度学习框架以及相应的深度神经网络模型，学习与验证神经网络实现图像分割任务的基本原理，'
         '掌握构建FCN、DeepLab等典型网络实现图像分割任务的基本方法，'
         '从而实现能够实现复杂分割任务的设计目标。', indent=True)

    # ===== 二、实验原理 =====
    heading(doc, '二、实验原理', 1)

    heading(doc, '2.1 图像分割任务概述', 2)
    para(doc,
         '图像分割（Semantic Segmentation）是计算机视觉中的基础任务，'
         '目标是对图像中每个像素进行分类，将其分配到预定义的语义类别中。'
         '与目标检测输出边界框不同，分割输出与原图等尺寸的像素级标签图（Label Map）。'
         '主要方法包括基于FCN的全卷积方法、基于编解码器结构的方法（如U-Net、SegNet）、'
         '以及引入空洞卷积和条件随机场的方法（如DeepLab系列）。', indent=True)

    heading(doc, '2.2 FCN（全卷积网络）', 2)
    para(doc,
         'FCN由Long等（2015）提出，是将CNN应用于语义分割的开创性工作。其核心思想是：', indent=True)
    para(doc,
         '(1) 全卷积化：将传统CNN中的全连接层替换为卷积层，使网络可以接受任意尺寸的输入；', indent=True)
    para(doc,
         '(2) 上采样（Upsampling）：使用转置卷积（反卷积）将低分辨率的特征图恢复到原始图像尺寸；', indent=True)
    para(doc,
         '(3) 跳跃连接（Skip Connection）：融合深层的语义信息和浅层的空间细节信息，'
         'FCN-8s融合了第3、4、5层的特征，比FCN-32s（仅使用第5层）具有更精细的分割边界。', indent=True)
    para(doc,
         'FCN-ResNet50使用ResNet-50作为骨干网络，通过空洞卷积保持特征图分辨率，'
         '最终通过双线性插值上采样到输入尺寸。', indent=True)

    heading(doc, '2.3 DeepLabV3', 2)
    para(doc,
         'DeepLabV3由Chen等（2017）提出，在FCN基础上引入了多项改进：', indent=True)
    para(doc,
         '(1) 空洞空间金字塔池化（ASPP）：使用不同膨胀率（dilation rate）的空洞卷积并行提取多尺度特征，'
         '捕捉不同大小的目标。同时加入全局平均池化分支获取全局上下文信息；', indent=True)
    para(doc,
         '(2) 空洞卷积（Atrous/Dilated Convolution）：在不降低分辨率的情况下扩大感受野，'
         '膨胀率r=1为标准卷积，r=2时隔一个像素采样，有效感受野扩大但参数量不变；', indent=True)
    para(doc,
         '(3) 改进的骨干网络：使用ResNet-50并在最后两个block中使用空洞卷积，'
         '保持高分辨率特征图；', indent=True)
    para(doc,
         '(4) 条件随机场（CRF）后处理（DeepLabV1/V2使用，V3中移除）：'
         '利用像素间的颜色相似性和空间接近性细化分割边界。', indent=True)

    heading(doc, '2.4 上采样方法', 2)
    para(doc,
         '语义分割需要将低分辨率特征图恢复到输入尺寸，常用的上采样方法包括：', indent=True)
    para(doc,
         '(1) 转置卷积（Transposed Convolution）：可学习的上采样，通过反卷积操作恢复空间分辨率；', indent=True)
    para(doc,
         '(2) 双线性插值（Bilinear Interpolation）：固定权重的上采样，计算量小但不可学习；', indent=True)
    para(doc,
         '(3) 像素洗牌（Pixel Shuffle）：通过通道到空间的变换实现上采样。', indent=True)

    # ===== 三、实验环境与数据集 =====
    heading(doc, '三、实验环境与数据集', 1)

    heading(doc, '3.1 实验环境', 2)
    para(doc, '编程语言：Python 3.9', indent=True)
    para(doc, '深度学习框架：PyTorch 2.7.0, torchvision 0.22.0', indent=True)
    para(doc, '运行设备：CPU', indent=True)
    para(doc, '依赖库：numpy, matplotlib, opencv, Pillow', indent=True)

    heading(doc, '3.2 测试数据', 2)
    para(doc,
         '本实验使用真实场景图像进行分割演示，图像来自日常自然场景（如室内、街道、'
         '室外环境等）。预训练模型在Pascal VOC数据集上训练，支持21类语义分割（含背景类）。'
         '相比简单几何合成图，真实图像更能体现模型对实际语义区域的划分能力。', indent=True)
    img(doc, os.path.join(IMG_DIR, 'test_images.png'), '图1 测试图像')

    # ===== 四、模型构建 =====
    heading(doc, '四、模型构建', 1)

    heading(doc, '4.1 FCN-ResNet50', 2)
    para(doc,
         '使用torchvision提供的预训练FCN-ResNet50模型，骨干网络为ResNet-50。'
         '模型在Pascal VOC数据集上预训练，支持21类语义分割。', indent=True)

    heading(doc, '4.2 DeepLabV3-ResNet50', 2)
    para(doc,
         '使用torchvision提供的预训练DeepLabV3-ResNet50模型，包含ASPP模块。'
         '模型在Pascal VOC数据集上预训练，支持21类语义分割。', indent=True)

    # ===== 五、实验结果与分析 =====
    heading(doc, '五、实验结果与分析', 1)

    heading(doc, '5.1 分割结果对比', 2)
    img(doc, os.path.join(IMG_DIR, 'segmentation_comparison.png'), '图2 FCN vs DeepLabV3 分割结果对比')

    para(doc,
         '从分割结果可视化可以看出，两种模型都能对图像中的主要区域进行大致的语义分类。'
         'DeepLabV3由于引入了ASPP模块和空洞卷积，在边缘细节和多尺度目标的处理上表现更优。'
         'FCN的分割结果相对粗糙，边界不够平滑。', indent=True)

    heading(doc, '5.2 详细分割结果', 2)
    for i in range(1, 7):
        detail_path = os.path.join(IMG_DIR, f'segmentation_detail_{i}.png')
        if os.path.exists(detail_path):
            img(doc, detail_path, f'图{2+i} 测试图{i}分割详细结果')

    heading(doc, '5.3 模型参数与推理速度对比', 2)
    img(doc, os.path.join(IMG_DIR, 'model_comparison.png'), '图9 模型参数量、推理速度与类别数对比')
    img(doc, os.path.join(IMG_DIR, 'pixel_agreement.png'), '图10 FCN与DeepLabV3分割结果像素一致性')

    para(doc,
         '从模型对比数据可以看出，DeepLabV3由于包含ASPP等额外模块，参数量略高于FCN。'
         '在CPU推理速度方面，两者差异取决于输入图像尺寸和模型复杂度。'
         '像素一致性分析显示了两种模型在分割结果上的吻合程度，'
         '较高的一致性说明两者在主要语义区域的判断基本一致，差异主要在边界细节。', indent=True)

    heading(doc, '5.4 结构差异分析', 2)
    para(doc,
         'FCN与DeepLabV3的核心结构差异如下：', indent=True)
    para(doc,
         '(1) 多尺度处理：FCN通过跳跃连接融合多尺度特征，DeepLabV3通过ASPP并行提取多尺度特征，'
         '后者在处理不同大小目标时更加灵活；', indent=True)
    para(doc,
         '(2) 感受野：DeepLabV3使用空洞卷积扩大感受野，不增加参数量也不降低分辨率，'
         '能捕捉更大范围的上下文信息；', indent=True)
    para(doc,
         '(3) 全局信息：DeepLabV3的ASPP包含全局平均池化分支，显式引入全局上下文信息，'
         '有助于理解图像整体语义；', indent=True)
    para(doc,
         '(4) 边界质量：DeepLabV3通常产生更平滑、更精确的分割边界，'
         '这得益于其更强大的特征提取能力。', indent=True)

    # ===== 六、总结与心得体会 =====
    heading(doc, '六、总结与心得体会', 1)
    para(doc,
         '通过本次图像分割实验，我深入理解了FCN和DeepLabV3两种典型分割网络的设计理念，'
         '主要收获如下：', indent=True)
    para(doc,
         '1. 全卷积思想：FCN将传统分类网络改造为像素级预测网络的核心思想——'
         '去除全连接层、使用转置卷积上采样——为后续所有分割网络奠定了基础。', indent=True)
    para(doc,
         '2. 多尺度特征的重要性：无论是FCN的跳跃连接还是DeepLabV3的ASPP，'
         '都强调了多尺度特征融合对分割精度的关键作用。深层特征提供语义信息，'
         '浅层特征提供空间细节。', indent=True)
    para(doc,
         '3. 空洞卷积的巧妙设计：在不增加参数量的情况下扩大感受野，'
         '是解决"分辨率 vs 感受野"矛盾的有效方法。', indent=True)
    para(doc,
         '4. 分割与检测的关系：分割提供了比检测更精细的像素级信息，'
         '在自动驾驶（道路/行人分割）、医学影像（器官/病灶分割）等领域有广泛应用。', indent=True)
    para(doc,
         '5. 预训练模型的价值：通过迁移学习，预训练的分割模型可以快速应用于新场景，'
         '大大降低了对标注数据量的需求。', indent=True)

    # ===== 附录 =====
    heading(doc, '附录：实验代码', 1)
    with open(CODE_FILE, 'r', encoding='utf-8') as f:
        code_block(doc, f.read())

    doc.save(OUTPUT_PATH)
    print(f'报告已生成: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
